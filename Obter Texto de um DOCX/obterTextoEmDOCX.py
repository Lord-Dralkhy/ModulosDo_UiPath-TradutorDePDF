"""
# Obtendo Texto de um Documento DOCX v1.2.0

## Visão Geral

Este script em Python utiliza a biblioteca `tkinter` para criar uma interface gráfica simples que solicita ao usuário o caminho de um arquivo DOCX.
O programa, em seguida, lê o conteúdo do arquivo e exibe cada parágrafo em caixas de diálogo separadas.

## Pré-requisitos

Certifique-se de ter a biblioteca `tkinter` e `python-docx` instaladas. Você pode instalar essas dependências usando o seguinte comando:

pip install tk python-docx

## Como Executar

Execute o script e uma janela de diálogo será exibida solicitando o caminho do arquivo DOCX. Insira o caminho ou clique em "Cancelar" para encerrar o programa.

## Aviso

- Este programa foi criado por Lord Dralkhy - Magno da Silva Gomes.
- GitHub: https://github.com/Lord-Dralkhy

Aproveite a simplicidade e eficiência para obter texto de um documento DOCX!
"""

import tkinter as tk

from docx import Document
from tkinter import simpledialog, messagebox

#------------------------------------------

def obterCaminhoDocx():

    # Dando um feedback para o usuário
    print('Solicitando arquivo')

    #Solicita e retorna o caminho do arquivo DOCX
    return simpledialog.askstring("Escolher DOCX", "Digite o caminho do arquivo DOCX (com o formato):")

def main():

    # Dando um feedback para o usuário
    print('Iniciando sistema')

    # Cria uma janela principal oculta
    root = tk.Tk()
    root.withdraw()

    # Solicita o caminho do arquivo DOCX
    DOCXSelecionado = obterCaminhoDocx()

    #------------------------------------------

    if not DOCXSelecionado:

        # Se o usuário cancelar, encerre o programa
        messagebox.showinfo("Aviso", "Operação cancelada. Encerrando o programa.")

    else:

        try:

            # Lê o documento DOCX
            doc = Document(DOCXSelecionado)

            totalParagrafos = str(len(doc.paragraphs))

            # Exibe o total de parágrafos no documento
            messagebox.showinfo("Total de parágrafos", totalParagrafos)

            # Exibe cada parágrafo
            for i, paragraph in enumerate(doc.paragraphs, start=0):

                # Exibe o indice do parágrafo selecionado, para fins de controle
                messagebox.showinfo("Index", str(i))

                # Exibe o texto
                messagebox.showinfo(f"Texto {str(i + 1)}/{totalParagrafos}", paragraph.text)

        except Exception as e:

            # Em caso de erro, exibe uma mensagem
            messagebox.showerror("Erro", f"Erro ao processar o arquivo: {str(e)}")

    # Dando um feedback para o usuário
    print('Feito')

if __name__ == "__main__":
    main()
