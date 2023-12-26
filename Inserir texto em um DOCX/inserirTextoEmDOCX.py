'''
# Inserção de Texto em um Documento DOCX v1.0.0

Este é um programa simples que permite ao usuário inserir texto em um parágrafo específico de um arquivo DOCX usando a biblioteca `python-docx`.

O código está disponível no GitHub: https://github.com/Lord-Dralkhy/ModulosDo_UiPath-TradutorDePDF/tree/main/Inserir%20texto%20em%20um%20DOCX

## Instruções de Uso:

1. Execute o programa e uma interface gráfica será exibida.
2. Insira o caminho do arquivo DOCX quando solicitado, incluindo o formato (.docx).
3. Insira o índice do parágrafo onde deseja inserir o texto.
4. Digite o texto que deseja adicionar ao parágrafo.
5. Clique em "OK" para concluir a inserção do texto.

## Observação:

- Se ocorrerem erros durante a execução, o programa tentará fornecer informações detalhadas na mensagem de erro resultante.

## Desenvolvido por:

Lord Dralkhy - Magno da Silva Gomes
GitHub: https://github.com/Lord-Dralkhy

Aproveite a simplicidade e eficiência para inserir texto em um documento DOCX!
'''

import tkinter as tk

from docx import Document
from docx.shared import Pt, RGBColor
from tkinter import simpledialog, messagebox

# Constantes para tipos de mensagem
TIPO_INFO = "info"
TIPO_ERRO = "error"

#------------------------------------------

def obterCaminhoDocx():

    # Dando um feedback para o usuário
    print('Solicitando arquivo')

    #Solicita e retorna o caminho do arquivo DOCX
    return simpledialog.askstring("Escolher DOCX", "Digite o caminho do arquivo DOCX (com o formato):")

#---

def exibirMensagem(titulo, mensagem, tipo):

    # Exibe uma mensagem na tela com um botão "OK"
    if tipo == TIPO_INFO:

        messagebox.showinfo(titulo, mensagem)

    elif tipo == TIPO_ERRO:

        messagebox.showerror(titulo, mensagem)

#---

def inserirTextoEmParagrafo(DOCXSelecionado):

    try:

        # Lê o documento DOCX
        doc = Document(DOCXSelecionado)

        # Pedir o indice do parágrafo no documento
        paragrafoIndex = int(simpledialog.askstring("Inserir parágrafo", "Digite o indice do parágrafo:"))

        # Pedir o texto do parágrafo no documento
        paragrafoTexto = simpledialog.askstring("Inserir texto", "Digite o texto para inserir no parágrafo:")

        # Verifica se o índice fornecido é válido
        if 0 <= paragrafoIndex < len(doc.paragraphs):

            # Obtém o parágrafo pelo índice
            paragrafo = doc.paragraphs[paragrafoIndex]

            # Remove todo o conteúdo do parágrafo
            for run in paragrafo.runs:
                run.clear()

            # Adiciona o novo texto mantendo a formatação original
            run = paragrafo.add_run(paragrafoTexto)

            # Copia a formatação original do parágrafo para o novo texto
            run.bold = paragrafo.runs[0].bold
            run.italic = paragrafo.runs[0].italic
            run.underline = paragrafo.runs[0].underline
            run.font.size = Pt(paragrafo.runs[0].font.size.pt)
            run.font.name = paragrafo.runs[0].font.name
            run.font.color.rgb = paragrafo.runs[0].font.color.rgb
            run.font.all_caps = paragrafo.runs[0].font.all_caps
            run.font.small_caps = paragrafo.runs[0].font.small_caps
            run.font.double_strike = paragrafo.runs[0].font.double_strike
            run.font.subscript = paragrafo.runs[0].font.subscript
            run.font.superscript = paragrafo.runs[0].font.superscript
            run.font.underline = paragrafo.runs[0].font.underline
            run.font.color.theme_color = paragrafo.runs[0].font.color.theme_color
            run.font.highlight_color = paragrafo.runs[0].font.highlight_color
            run.font.shadow = paragrafo.runs[0].font.shadow
            run.font.outline = paragrafo.runs[0].font.outline

            # Salva as alterações no documento
            doc.save(DOCXSelecionado)

            print(f"Texto inserido com sucesso no parágrafo {paragrafoIndex}.")

        else:

            exibirMensagem("Aviso", f"Índice {paragrafoIndex} inválido. O documento possui {len(doc.paragraphs)} parágrafos.", "error")

            print(f"Índice {paragrafoIndex} inválido. O documento possui {len(doc.paragraphs)} parágrafos.")

            return "erro"

        return "finalizar"

    except Exception as e:

        # Em caso de erro, exibe uma mensagem
        messagebox.showerror("Erro", f"Erro ao processar o arquivo: {str(e)}")

        print(f"Erro ao processar o arquivo: {str(e)}")

        return "erro"

#---

def encerrarPrograma():

    # Se o usuário cancelar, encerre o programa
    exibirMensagem("Aviso", "Operação cancelada. Encerrando o programa.", "error")

    # Dando um feedback para o usuário
    exibirMensagem("Criado por:", "Este programa foi feito por:\nLord Dralkhy - Magno da Silva Gomes", "info")

#------------------------------------------

def main():

    # Dando um feedback para o usuário
    print("Iniciando sistema")

    # Cria uma janela principal oculta
    root = tk.Tk()
    root.withdraw()

    #---

    progresso = "erro"

    while progresso == "erro":

        DOCXSelecionado = ""

        # Verificar se a variável está vazia (que é diferente de ser nula)
        while DOCXSelecionado == "":

            # Solicita o caminho do arquivo DOCX
            DOCXSelecionado = obterCaminhoDocx()

        # ---

        # Se ela for nula, cancelar o programa
        if DOCXSelecionado is None:

            encerrarPrograma()

            progresso = "finalizar"

        else:

            progresso = inserirTextoEmParagrafo(DOCXSelecionado)

# ------------------------------------------

    # Dando um feedback para o usuário
    print('Feito')

if __name__ == "__main__":
    main()